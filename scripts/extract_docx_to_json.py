import os
import re
import json
import hashlib
import csv
from docx import Document
from docx.oxml.ns import qn
from PIL import Image
import pytesseract
import io

# Configurer le chemin de Tesseract
pytesseract.pytesseract.tesseract_cmd = r'E:\Tesseract-OCR\tesseract.exe'
# Configurer TESSDATA_PREFIX
os.environ['TESSDATA_PREFIX'] = r'E:\Tesseract-OCR\tessdata'

def normalize_ws(s: str) -> str:
    """Normalise les espaces Unicode en espaces réguliers et supprime les caractères invisibles."""
    return (
        s.replace("\u00A0", " ")  # NBSP
         .replace("\u202F", " ")  # NARROW NBSP
         .replace("\u2009", " ")  # THIN SPACE
         .replace("\u200A", " ")  # HAIR SPACE
         .replace("\u2007", " ")  # FIGURE SPACE
         .replace("\u2002", " ")  # EN SPACE
         .replace("\u2003", " ")  # EM SPACE
         .replace("\u2004", " ")  # THREE-PER-EM SPACE
         .replace("\u2005", " ")  # FOUR-PER-EM SPACE
         .replace("\u2006", " ")  # SIX-PER-EM SPACE
         .replace("\u2008", " ")  # PUNCTUATION SPACE
         .replace("\u205F", " ")  # MEDIUM MATHEMATICAL SPACE
         .replace("\u200B", "")   # ZERO WIDTH SPACE
         .replace("\ufeff", "")   # BOM
    ).strip()

def is_numeric_heading(title: str) -> bool:
    """Retourne True si le titre est un code numérique (ex. : '1', '1.2', '2.10.3')."""
    return bool(re.fullmatch(r"\d+(?:\.\d+)*", title))

def get_heading_code(title: str):
    """Extrait le code numérique initial (ex. : '2.0', '2.18') d'un titre, sinon None."""
    m = re.match(r"^(\d+(?:\.\d+)*)\b", title.strip())
    return m.group(1) if m else None

def is_fragment_line(line: str) -> bool:
    """Retourne True si la ligne contient 1-3 lettres (ex. : 'g', 'r e')."""
    s = normalize_ws(line)
    s = re.sub(r"\s+", "", s)
    letters = re.sub(r"[^A-Za-zÀ-ÿ]", "", s)
    return 1 <= len(letters) <= 3 and bool(re.fullmatch(r"[A-Za-zÀ-ÿ]+", letters))

def merge_fragment_into_title(title: str, fragment: str) -> str:
    """Ajoute les lettres du fragment à la fin du dernier mot du titre."""
    tit = normalize_ws(title)
    frag = normalize_ws(fragment)
    frag = re.sub(r"\s+", "", frag)
    if not frag:
        return tit
    return (tit + frag) if re.search(r"[A-Za-zÀ-ÿ]$", tit) else (tit + " " + frag)

def extract_images_from_doc(doc, images_dir: str, image_hashes: set, image_by_hash: dict):
    """Extrait les images du document et les sauvegarde dans images_dir. Retourne une liste de métadonnées."""
    os.makedirs(images_dir, exist_ok=True)
    out = []
    for i, rel in enumerate(doc.part.rels.values()):
        if "image" in rel.target_ref:
            image_bytes = rel.target_part.blob
            img_hash = hashlib.md5(image_bytes).hexdigest()
            if img_hash in image_hashes:
                out.append(image_by_hash[img_hash])
                continue
            image_hashes.add(img_hash)
            pil_img = Image.open(io.BytesIO(image_bytes))
            if pil_img.mode != 'RGB':
                pil_img = pil_img.convert('RGB')
            filename = f"img_{i}_{img_hash[:8]}.png"
            filepath = os.path.join(images_dir, filename)
            pil_img.save(filepath, 'PNG')
            # Appliquer l'OCR avec gestion d'erreurs
            try:
                ocr_text = pytesseract.image_to_string(pil_img, lang='fra')
            except pytesseract.TesseractError as e:
                print(f"Erreur OCR pour l'image {filename}: {e}")
                ocr_text = ""
            meta = {
                'index': i,
                'filename': filename,
                'filepath': filepath,
                'url': f"/assets/images/{filename}",
                'hash': img_hash,
                'ocr_text': normalize_ws(ocr_text) if ocr_text else ""
            }
            image_by_hash[img_hash] = meta
            out.append(meta)
    return out

def extract_tables_from_doc(doc, tables_dir: str, table_hashes: set, table_by_hash: dict):
    """Extrait les tableaux du document et les sauvegarde en CSV. Retourne une liste de métadonnées."""
    os.makedirs(tables_dir, exist_ok=True)
    out = []
    for tbl_index, table in enumerate(doc.tables):
        # Convertir le tableau en liste de listes pour le hachage et CSV
        tbl_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        tbl_str = json.dumps(tbl_data, ensure_ascii=False)
        tbl_hash = hashlib.md5(tbl_str.encode('utf-8')).hexdigest()
        if tbl_hash in table_hashes:
            out.append(table_by_hash[tbl_hash])
            continue
        table_hashes.add(tbl_hash)
        filename = f"tbl_{tbl_index}_{tbl_hash[:8]}.csv"
        filepath = os.path.join(tables_dir, filename)
        with open(filepath, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerows(tbl_data)
        meta = {
            'index': tbl_index,
            'filename': filename,
            'filepath': filepath,
            'url': f"/assets/tables/{filename}",
            'hash': tbl_hash
        }
        table_by_hash[tbl_hash] = meta
        out.append(meta)
    return out

def parse_docx_to_json(docx_path, output_json):
    """Parse un fichier .docx et extrait la hiérarchie, images et tableaux dans une structure JSON."""
    doc = Document(docx_path)
    structure = []
    current_chapter = None
    current_section = None
    current_subsection = None
    sections_index = {}

    # Gestion des images
    images_dir = os.path.join('extracted_assets', 'images')
    image_hashes = set()
    image_by_hash = {}
    all_images = extract_images_from_doc(doc, images_dir, image_hashes, image_by_hash)
    image_ptr = 0

    # Gestion des tableaux
    tables_dir = os.path.join('extracted_assets', 'tables')
    table_hashes = set()
    table_by_hash = {}
    all_tables = extract_tables_from_doc(doc, tables_dir, table_hashes, table_by_hash)
    table_ptr = 0

    for para in doc.paragraphs:
        text = normalize_ws(para.text)
        if not text:
            continue

        # Détection des balises hiérarchiques via regex
        m_ch = re.match(r'^\d+\.0\s*(.*)$', text)  # Chapitre (ex: "1.0")
        m_sec = re.match(r'^\d+\.\d+\s*(.*)$', text)  # Section (ex: "1.1")
        m_sub = re.match(r'^\d+\.\d+\.\d+\s*(.*)$', text)  # Sous-section (ex: "1.1.1")

        if m_ch:  # Chapitre
            ch_title = normalize_ws(m_ch.group(0))
            existing_chapter = next((ch for ch in structure if ch.get("title") == ch_title), None)
            if existing_chapter:
                current_chapter = existing_chapter
            else:
                current_chapter = {
                    "title": ch_title,
                    "content": [],
                    "sections": [],
                    "images": [],
                    "tables": []
                }
                structure.append(current_chapter)
            current_section = None
            current_subsection = None
            sections_index = {get_heading_code(sec['title']): sec for sec in current_chapter['sections'] if get_heading_code(sec['title'])}

        elif m_sec:  # Section
            sec_title = normalize_ws(m_sec.group(0))
            code = get_heading_code(sec_title)
            existing_section = None
            for chapter in structure:
                for section in chapter.get("sections", []):
                    if section.get("title") == sec_title or (code and get_heading_code(section.get("title")) == code):
                        existing_section = section
                        break
                if existing_section:
                    break

            if existing_section:
                current_section = existing_section
                if len(sec_title) > len(current_section.get("title", "")):
                    current_section["title"] = sec_title
            else:
                if current_chapter is None:
                    current_chapter = {"title": "Chapitre", "content": [], "sections": [], "images": [], "tables": []}
                    structure.append(current_chapter)
                current_section = {
                    "title": sec_title,
                    "content": [],
                    "subsections": [],
                    "images": [],
                    "tables": []
                }
                current_chapter["sections"].append(current_section)
                if code:
                    sections_index[code] = current_section
            current_subsection = None

        elif m_sub:  # Sous-section
            sub_title = normalize_ws(m_sub.group(0))
            current_subsection = {
                "title": sub_title,
                "content": [],
                "images": [],
                "tables": []
            }
            if current_section is None:
                current_section = {
                    "title": "Section",
                    "content": [],
                    "subsections": [],
                    "images": [],
                    "tables": []
                }
                if current_chapter is None:
                    current_chapter = {"title": "Chapitre", "content": [], "sections": [], "images": [], "tables": []}
                    structure.append(current_chapter)
                current_chapter["sections"].append(current_section)
            current_section["subsections"].append(current_subsection)

        else:  # Contenu normal
            # Détection des légendes d'images (ex: "Fig.4.1 Manilledecelluledynamométrique")
            m_imgcap = re.match(r'^Fig\.(\d+\.\d+)\s*(.*)$', text, re.IGNORECASE)
            if m_imgcap and image_ptr < len(all_images):
                img_title = normalize_ws(m_imgcap.group(2)) if m_imgcap.group(2) else f"Image {m_imgcap.group(1)}"
                img_meta = all_images[image_ptr]
                image_ptr += 1
                target = current_subsection or current_section or current_chapter
                if target is not None:
                    target.setdefault("images", [])
                    target["images"].append({
                        "title": img_title,
                        "url": img_meta["url"],
                        "filename": img_meta["filename"],
                        "index": img_meta["index"],
                        "ocr_text": img_meta["ocr_text"]
                    })
                continue

            # Détection des légendes de tableaux (ex: "Tableau 1 Chargesderupturedescâblesmétalliques")
            m_tblcap = re.match(r'^Tableau\s+(\d+)\s*(.*)$', text, re.IGNORECASE)
            if m_tblcap and table_ptr < len(all_tables):
                num = m_tblcap.group(1)
                tbl_title = normalize_ws(m_tblcap.group(2)) if m_tblcap.group(2) else f"Tableau {num}"
                tbl_meta = all_tables[table_ptr]
                table_ptr += 1
                target = current_subsection or current_section or current_chapter
                if target is not None:
                    target.setdefault("tables", [])
                    target["tables"].append({
                        "title": tbl_title,
                        "url": tbl_meta["url"],
                        "filename": tbl_meta["filename"],
                        "index": tbl_meta["index"]
                    })
                continue

            # Gestion des titres numériques suivis de texte descriptif
            if current_subsection is not None and is_numeric_heading(current_subsection["title"]):
                if not re.match(r'^\s*\d+\.\d+', text):
                    new_title = normalize_ws(f"{current_subsection['title']} {text}")
                    for ss in list(current_section["subsections"]):
                        if ss is current_subsection:
                            continue
                        if ss["title"] == new_title:
                            ss["content"].extend(current_subsection.get("content", []))
                            current_section["subsections"].remove(current_subsection)
                            current_subsection = ss
                            break
                    current_subsection["title"] = new_title
                    continue

            if current_section is not None and is_numeric_heading(current_section["title"]):
                if not re.match(r'^\s*\d+\.\d+', text):
                    new_title = normalize_ws(f"{current_section['title']} {text}")
                    for sec in list(current_chapter["sections"]):
                        if sec is current_section:
                            continue
                        if sec["title"] == new_title:
                            sec["content"].extend(current_section.get("content", []))
                            sec["subsections"].extend(current_section.get("subsections", []))
                            current_chapter["sections"].remove(current_section)
                            current_section = sec
                            break
                    current_section["title"] = new_title
                    continue

            # Merge fragments into titles
            if current_subsection is not None and is_fragment_line(text):
                current_subsection["title"] = merge_fragment_into_title(current_subsection["title"], text)
                continue
            if current_section is not None and is_fragment_line(text):
                current_section["title"] = merge_fragment_into_title(current_section["title"], text)
                continue
            if current_chapter is not None and is_fragment_line(text):
                current_chapter["title"] = merge_fragment_into_title(current_chapter["title"], text)
                continue

            # Ajouter au contenu
            if current_subsection is not None:
                current_subsection["content"].append(text)
            elif current_section is not None:
                current_section["content"].append(text)
            elif current_chapter is not None:
                current_chapter["content"].append(text)

    # Sauvegarde JSON
    with open(output_json, "w", encoding="utf-8") as f:
        json.dump(structure, f, indent=4, ensure_ascii=False)

    return structure

if __name__ == "__main__":
    docx_path = r"D:\GitHub\digitalbook\scripts\Livre digital .docx"
    output_json = "structure.json"
    try:
        data = parse_docx_to_json(docx_path, output_json)
        print(json.dumps(data, indent=4, ensure_ascii=False))
    except Exception as e:
        print(f"Erreur lors de l'exécution : {e}")